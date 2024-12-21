from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pymongo import MongoClient
import io
import fitz  # PyMuPDF
import docx
import uuid

app = Flask(__name__)
CORS(app)

# MongoDB-Verbindung
client = MongoClient("mongodb://db:27017")
db = client.pdf_database  # Datenbank
pdf_collection = db.pdfs  # Collection für PDF-Daten

# Vorhandene PDF-Konvertierungsfunktionen
def convert_pdf_to_text(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def convert_pdf_to_word(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    word_doc = docx.Document()
    for page in doc:
        word_doc.add_paragraph(page.get_text())
    word_stream = io.BytesIO()
    word_doc.save(word_stream)
    word_stream.seek(0)
    return word_stream

def convert_pdf_to_html(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    html_content = "<html><body>"
    for page in doc:
        html_content += f"<h2>Page {page.number + 1}</h2><p>{page.get_text()}</p>"
    html_content += "</body></html>"
    html_stream = io.BytesIO(html_content.encode('utf-8'))
    html_stream.seek(0)
    return html_stream

# Neue Funktionen zur Konvertierung des Chatverlaufs
def convert_chat_to_text(chat_history):
    text = "\n".join([f"Du: {msg['user']}\nBot: {msg['bot']}" for msg in chat_history])
    return text

def convert_chat_to_word(chat_history):
    word_doc = docx.Document()
    for msg in chat_history:
        word_doc.add_paragraph(f"Du: {msg['user']}\nBot: {msg['bot']}")
    word_stream = io.BytesIO()
    word_doc.save(word_stream)
    word_stream.seek(0)
    return word_stream

def convert_chat_to_html(chat_history):
    html_content = "<html><body>"
    for msg in chat_history:
        html_content += f"<p><strong>Du:</strong> {msg['user']}<br><strong>Bot:</strong> {msg['bot']}</p>"
    html_content += "</body></html>"
    html_stream = io.BytesIO(html_content.encode('utf-8'))
    html_stream.seek(0)
    return html_stream

@app.route('/convert', methods=['POST'])
def convert_file():
    data = request.json
    file_id = data.get('file_id', '')
    conversion_type = data.get('conversion_type', 'text')

    if not file_id:
        return jsonify({'error': 'No file_id provided in the request'}), 400

    # PDF-Daten aus MongoDB laden
    pdf_data = pdf_collection.find_one({"_id": file_id})
    if not pdf_data:
        return jsonify({'error': 'PDF not found in the database'}), 404

    file_bytes = pdf_data.get("file_data", None)
    if not file_bytes:
        return jsonify({'error': 'No file data found in the database'}), 404

    # Konvertierung basierend auf dem Typ
    if conversion_type == 'text':
        text = convert_pdf_to_text(file_bytes)
        return send_file(io.BytesIO(text.encode('utf-8')), mimetype='text/plain', as_attachment=True, download_name=f'{pdf_data["filename"]}.txt')
    elif conversion_type == 'word':
        word_stream = convert_pdf_to_word(file_bytes)
        return send_file(word_stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f'{pdf_data["filename"]}.docx')
    elif conversion_type == 'html':
        html_stream = convert_pdf_to_html(file_bytes)
        return send_file(html_stream, mimetype='text/html', as_attachment=True, download_name=f'{pdf_data["filename"]}.html')
    else:
        return jsonify({'error': 'Invalid conversion type'}), 400

@app.route('/convert-chat', methods=['POST'])
def convert_chat():
    chat_history = request.json.get('chat_history', [])
    conversion_type = request.json.get('conversion_type', 'text')

    if not chat_history:
        return jsonify({'error': 'No chat history provided'}), 400

    if conversion_type == 'text':
        text = convert_chat_to_text(chat_history)
        return send_file(io.BytesIO(text.encode('utf-8')), mimetype='text/plain', as_attachment=True, download_name='chat_history.txt')
    elif conversion_type == 'word':
        word_stream = convert_chat_to_word(chat_history)
        return send_file(word_stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='chat_history.docx')
    elif conversion_type == 'html':
        html_stream = convert_chat_to_html(chat_history)
        return send_file(html_stream, mimetype='text/html', as_attachment=True, download_name='chat_history.html')
    else:
        return jsonify({'error': 'Invalid conversion type'}), 400

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5003)
